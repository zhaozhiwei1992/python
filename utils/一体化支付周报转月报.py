# !user/bin/python
# _*_ coding: utf-8 _*_
#
# @Title: 一体化支付周报转月报.py
# @Description: 每个月需要整理月报, 将周报合并到月报中
# 1. 需要存在月报模板excel
# 2. 整理好每个月所属的周报
# @author zhaozhiwei
# @date 2022/10/31 下午2:52
# @version V1.0


import openpyxl
import docx
import os
import datetime


def getWeekDocDataList(year, month):
    """
    遍历获取当月周报信息

    构建周报数据, 返回list<tuple>信息
    """

    # 列表信息
    tableColList = []

    # 1. 便利目录下当月周报, 文件名格式如: 周报-支付20220815-20220819.doc, 获取前缀是周报-支付202208*.doc即可
    # for filename in os.listdir("/mnt/d/00共享/04工作计划/01周计划/2022年"):
    for filename in os.listdir("/mnt/d/00共享/04工作计划/01周计划/{}年/支付周报/".format(year)):
        # 绝对路径
        fileFullName = "/mnt/d/00共享/04工作计划/01周计划/{}年/支付周报/{}".format(year, filename)
        if filename.startswith("周报-支付{}{}".format(year, month)):
            # 2. 读取每个word数据, 进行构建
            file = docx.Document(fileFullName)
            # 读取所有的表格, 只保留跟上述表名能对应的
            # tableCount = len(file.tables)
            # print("有 ", tableCount, "个表, 取第二个")
            # 一个周报两个表格, 只取第二个本周周报
            curWeekTable = file.tables[1]
            rowCount = len(curWeekTable.rows)

            for i_row in range(1, rowCount):
                # 遍历每一个表格的每一行
                # doc.tables[i].rows[i_row]
                rowData = curWeekTable.rows[i_row]

                # 列信息
                #  0:序号(禅道号) 1:产品(区分业务系统) 2:项目(空则全国) 3:任务内容 4:责任人 5:预计工时(小时, 8的倍数) 计划完成时间 实际完成时间 完成百分比 偏差说明

                colRow = (rowData.cells[0].text, rowData.cells[1].text, rowData.cells[2].text
                          , rowData.cells[3].text
                          , rowData.cells[4].text, rowData.cells[5].text)
                tableColList.append(colRow)
    return tableColList


def weekDataWriteToMonthDoc(year, month, tableColList):
    """
    使用周报的数据生成为月报
    """
    # 1. 数据根据 1:产品(区分业务系统) 分组,
    groupDict = {}
    for i in tableColList:
        groupName = i[1]
        if len(str(groupName)) < 1:
            groupName = "其它"
        if groupName in groupDict.keys():
            tableCols = groupDict[groupName]
            tableCols.append(i)
        else:
            tableCols = [i]
            groupDict[groupName] = tableCols

    # print(groupDict)

    # 2. 每个分组分别写入不同的sheet页
    wb = openpyxl.load_workbook(
        '/home/zhaozhiwei/workspace/项目管理/一大堆模板/excel模板/月报-预算执行工作量统计模板.xlsx')
    # 分组key为系统名称, 跟sheet页匹配
    for appName in groupDict:
        sheet = wb[appName]
        #       A:禅道号	B:系统	C:需求及BUG问题描述 D:问题类型 E:所属项目	F:责任人	G:开发工作量（人/日）H:测试工作量（人/日）
        data = groupDict[appName]
        for index, ele in enumerate(data):
            #  0:序号(禅道号) 1:产品(区分业务系统) 2:项目(空则全国) 3:任务内容 4:责任人 5:预计工时(小时, 8的倍数) 计划完成时间 实际完成时间 完成百分比 偏差说明
            sheet['A' + str(index + 2)].value = ele[0]
            sheet['B' + str(index + 2)].value = ele[1]
            sheet['C' + str(index + 2)].value = ele[3]
            sheet['D' + str(index + 2)].value = "bug" if str(ele[0]).startswith("#") else "需求"
            sheet['E' + str(index + 2)].value = ele[2]
            sheet['F' + str(index + 2)].value = ele[4]
            sheet['G' + str(index + 2)].value = int(ele[5]) / 8
            sheet['H' + str(index + 2)].value = int(ele[5]) / 8

    wb.save('/tmp/预算一体化预算部分开发任务及工作量统计({}年{}月)-预算执行.xlsx'.format(year, month))


if __name__ == '__main__':

    now = datetime.date.today()
    # 获取年 月
    year = now.year
    month = now.month

    # 1. 遍历获取当月周报数据, 并写入到一个python对象中
    tableColList = getWeekDocDataList(str(year), str(month).zfill(2))
    # print(tableColList)
    # 2. 根据python对象分页签写入到月报excel中
    if len(tableColList) > 0:
        weekDataWriteToMonthDoc(str(year), str(month).zfill(2), tableColList)
    # 3. 手动编辑文档, 填充公式并计算
