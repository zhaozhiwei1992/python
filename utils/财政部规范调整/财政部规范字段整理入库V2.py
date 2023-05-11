# !user/bin/python
# _*_ coding: utf-8 _*_
#
# @Title: 财政部规范字段整理入库V2.py
# @Description:
# 将财政部规范excel中的表和字段信息导出到数据库
# @author zhaozhiwei
# @date 2022/8/4 上午11:15
# @version V1.0

import os

import cx_Oracle
import docx

import 财政部规范字段比对导出excel as exportExcel

# 设置查询编码
os.environ['NLS_LANG'] = 'AMERICAN_AMERICA.ZHS16GBK'

# 数据库配置
ip = '192.168.100.80'
port = 1521
SID = 'ORCL'
dsn = cx_Oracle.makedsn(ip, port, SID)

def getDocDataListV2(fileName="v2.0 .194-20220726预算管理一体化系统技术标准V2.0（中册）（征求意见稿）.docx"):
    # 读取word数据, 如何1表名对一个表格
    file = docx.Document("/tmp/" + fileName)
    print("总段落数:" + str(len(file.paragraphs)))
    tableNameList = []
    # 找到开头是表名的行, 如 表名：PAY_XX
    for para in file.paragraphs:
        if str(para.text).startswith("表名："):
            # print(para.text)
            tableNameList.append(str(para.text).replace("表名：", ""))
    # 321个表名称
    print("标识表名:xx的个数", len(tableNameList))
    # print(tableNameList)
    # 读取所有的表格, 只保留跟上述表名能对应的
    # 347
    tableCount = len(file.tables)
    print("文档中所有表格数: ", tableCount)
    # 列表信息
    tableColList = []
    tableCount2 = 0
    for i in range(0, tableCount):
        # doc.tables[i]                                         ## 遍历每一个表格
        rowCount = len(file.tables[i].rows)
        # 如果第一行第二列是  "字段名称" 才进行便利
        cellContent = file.tables[i].rows[0].cells[1].text
        # print(cellContent)

        # 321个表，与总表数匹配
        if str(cellContent).startswith("字段名称"):
            tableCount2 += 1
            for i_row in range(1, rowCount):
                # 遍历每一个表格的每一行
                rowData = file.tables[i].rows[i_row]
                colCount = len(rowData.cells)
                # print("colCount", colCount)

                # 跳过有问题的行
                if(colCount < 4 or colCount < 5 or colCount < 6 or colCount < 7):
                    print("表格有问题: ", rowData, "表名", tableNameList[tableCount2 - 1], "游标", tableCount2)
                    continue

                # 将文档中字段类型拆分
                if(None == rowData.cells[3]):
                    colTypeStr = "null"
                else:
                    colTypeStr = str(rowData.cells[3].text)
                colType = ""
                colLength = ""
                if "(" in colTypeStr:
                    # 英文括号
                    colType = colTypeStr.split("(")[0]
                    colLength = colTypeStr.split("(")[1].replace(")", "")
                elif "（" in colTypeStr:
                    # 中文括号
                    colType = colTypeStr.split("（")[0]
                    colLength = colTypeStr.split("（")[1].replace("）", "")
                else:
                    colType = colTypeStr
                    colLength = "0"

                # 列信息
                # 序号 字段名称 中文名称 类型及长度 强制 / 可选(M) 是否必填(是/否) 库表要素编号 备注
                # map形式数据
                # colMap = {"table_name": tableNameList[tableCount2-1], "col_code": rowData.cells[1].text,
                #           "col_name": rowData.cells[2].text, "type": colType,
                #           "length": colLength, "required": rowData.cells[4].text}

                colRow = (tableNameList[tableCount2 - 1], rowData.cells[1].text, rowData.cells[2].text
                          , colType, colLength, rowData.cells[4].text, rowData.cells[5].text)
                # 游标11匹配
                tableColList.append(colRow)
    # 321个表
    # 实际个数应与 表名:xx个数匹配
    print("实际使用表个数", tableCount2)
    # print("表列信息", tableColList)
    return tableColList

def saveDataToOracle(tableName, tableColList):
    """
    入库
    根据表名，及表列配置信息, 批量存储配置
    程序主动提交
    """
    # 连接数据库
    connection = cx_Oracle.connect('PAY_34', '1', dsn)
    try:
        # 测试连接用——输出数据库版本
        print(connection.version)
        # 获取游标
        cursor = cx_Oracle.Cursor(connection)
        # 写入操作
        cursor.prepare('insert into ' + tableName + ' (table_name, col_code, col_name, type, length, mo, required) '
                                                    'values '
                                                    '(:1, :2, :3, :4, :5, :6, :7)')
        # 执行入库
        cursor.executemany(None, tableColList)
        connection.commit()
        cursor.close()
        connection.close()

    except Exception as e:
        print('Oracle 写入失败，Exception:{0}'.format(e))
        connection.rollback()
        connection.close()

def getDocDataListV1(fileName="v2.0 .149-20210802预算管理一体化系统技术标准（下册）(只留修正版）.docx"):
    """
    获取1.0版本的列表信息
    """
    # 读取word数据, 如何1表名对一个表格
    file = docx.Document("/tmp/" + fileName)
    print("总段落数:" + str(len(file.paragraphs)))
    tableNameList = []
    # 找到开头是表名的行, 如 表名：PAY_XX
    for para in file.paragraphs:
        if str(para.text).startswith("表名："):
            # print(para.text)
            tableNameList.append(str(para.text).replace("表名：", ""))
    print(len(tableNameList))
    # 读取所有的表格, 只保留跟上述表名能对应的
    tableCount = len(file.tables)
    print(tableCount)
    # 列表信息
    tableColList = []
    tableCount2 = 0
    for i in range(0, tableCount):
        # doc.tables[i]                                         ## 遍历每一个表格
        rowCount = len(file.tables[i].rows)
        # 如果第一行第二列是  "字段名称" 才进行便利
        cellContent = file.tables[i].rows[0].cells[1].text
        # print(cellContent)

        # 321个表，与总表数匹配
        if str(cellContent).startswith("字段名称"):
            tableCount2 += 1
            for i_row in range(1, rowCount):
                # doc.tables[i].rows[i_row]                        ## 遍历每一个表格的每一行
                rowData = file.tables[i].rows[i_row]

                # 列信息
                # 序号 字段名称 中文名称 类型 长度 强制 / 可选 库表要素编号 备注
                # 0    1      2      3    4   5          6         7

                colRow = (tableNameList[tableCount2 - 1], rowData.cells[1].text, rowData.cells[2].text
                          , rowData.cells[3].text
                          , rowData.cells[4].text, rowData.cells[5].text, "是" if rowData.cells[5].text == "M" else "否")
                # 游标11匹配
                tableColList.append(colRow)
    # 178
    # tableCount2 == tableNameList.length  必须相等
    print(tableCount2)
    print("149编号表列信息", tableColList)

    return tableColList


if __name__ == '__main__':

    # 将读取数据整理标准格式, 方便写入数据库或者excel中
    # tableColList = getDocDataListV2("2.中册-预算管理一体化系统技术标准20230505.docx")
    # saveDataToOracle("STANDARD_FIELD_V2_20230505", tableColList)

    # tableColList = getDocDataListV1("预算管理一体化系统技术标准（下册）.docx")
    # saveDataToOracle("STANDARD_FIELD_V1_STANDARD", tableColList)

    # 数据生成后自行将表中table_name_cn进行填充, 调整一些有误的内容

    # 生成字段比对报告, 根据两个版本表比对
    exportExcel.compareToExcel("STANDARD_FIELD_V1_JC20230505", "STANDARD_FIELD_V2_20230505")
